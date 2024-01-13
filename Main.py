#Toda vez que inicializar o código pelo google colab é necessário instalar as bibliotecas
!pip install xlwings unidecode python-docx

import os
from unidecode import unidecode
from docx import Document
from google.colab import files
import glob

#Importando as função
from funcao_orientacao import orientacao
from funcao_processo_adm import processo_adm
from funcao_processo_adm_anexado import processo_adm_anexado
from funcao_prestadora import prestadora
from funcao_cnpj import cnpj
from funcao_servico import servico
from funcao_orgao_judicializado import orgao_judicializado
from funcao_vara import vara
from funcao_processo_judicial import processo_judicial
from funcao_tipo_processo import tipo_processo
from funcao_data_atuacao_judiciario import data_atuacao_judiciario
from funcao_data_decisao_liminar import data_decisao_liminar
from funcao_doc_comunicacao_judicializacao import doc_comunicacao_judicializacao
from funcao_status_processo_adm import status_processo_adm
from funcao_status_processo_judicial import status_processo_judicial
from funcao_data_ultima_consulta import data_ultima_consulta
from funcao_regulamento import regulamento
from funcao_palavras_chave import palavras_chave
from funcao_questionamento import questionamentos
from funcao_planilha import planilha
from funcao_data_1_instancia import data_1_instancia
from funcao_inclinacao_1 import inclinacao_1

print("Lendo os processos de um diretório específico")
pasta = '/content/'
lista_processos = []
arquivos = glob.glob(os.path.join(pasta, '*.docx'))
arquivos.sort(key = lambda filename: int(os.path.basename(filename).split('-')[0]))
for arquivo in arquivos:
  doc = Document(arquivo)
  lista_processos.append(arquivo)
print("Processos salvos")

#Uma lista para cada informação do processo que precisa ir para a planilha
numero_processo = []
processo_anexado = []
lista_prestadora = []
lista_cnpj = []
lista_servico = []
lista_orgao = []
lista_vara = []
processo_principal = []
processo_associado = []
tipologia_processo = []
data_atuacao = []
data_liminar = []
data_primeira_instancia = []
primeira_inclinacao = []
doc_judicializacao = []
status_processo_administrativo = []
status_processo = []
data_consulta = []
keywords = []
titulo = []
questionamento = []
tese = []
descricao_infracao = []
regulamento_oficial = []
dispositivo_oficial = []

#Loop que percorre todos os processos que foram salvos na lista_processos
for i in range(len(lista_processos)):
    print("Início do código")
    doc = Document(lista_processos[i])
    print("Processo: ", lista_processos[i])
    orientacao(doc)

    #Excluir a primeira linha de cada processo
    print("Excluindo primeira linha do processo")
    t = doc.tables[0]
    t._element.remove(t.rows[0]._element)
    print("Primeira linha do processo excluída")

    cell = t.cell(1,0)
    numero_processo.append(processo_adm(cell))

    cell = t.cell(1, 1)
    processo_anexado.append(processo_adm_anexado(cell))

    cell = t.cell(1, 2)
    lista_prestadora.append(prestadora(cell))

    cell = t.cell(1, 3)
    lista_cnpj.append(cnpj(cell))

    cell = t.cell(1, 4)
    lista_servico.append(servico(cell))

    cell = t.cell(1, 5)
    lista_orgao.append(orgao_judicializado(cell))

    cell = t.cell(1, 6)
    lista_vara.append(vara(cell))

    cell = t.cell(1, 7)
    cell.text, processo_join = processo_judicial(cell)
    processo_principal.append(cell.text)
    processo_associado.append(processo_join)

    cell = t.cell(1, 8)
    tipologia_processo.append(tipo_processo(cell))

    cell = t.cell(1, 9)
    data_atuacao.append(data_atuacao_judiciario(cell))

    cell = t.cell(1, 10)
    data_liminar.append(data_decisao_liminar(cell))

    cell = t.cell(1, 11)
    doc_judicializacao.append(doc_comunicacao_judicializacao(cell))

    cell = t.cell(1,14)
    data_primeira_instancia.append(data_1_instancia(cell))
    primeira_inclinacao.append(inclinacao_1(cell))

    cell = t.cell(1, 23)
    status_processo_administrativo.append(status_processo_adm(cell))

    cell = t.cell(1, 24)
    status_processo.append(status_processo_judicial(cell))

    cell = t.cell(1, 25)
    data_consulta.append(data_ultima_consulta(cell))

    cell = t.cell(1,20)
    lista_regulamentos, lista_dispositivos, descricao = regulamento(cell)
    descricao_infracao.append(descricao)
    regulamento_oficial.append(lista_regulamentos)
    dispositivo_oficial.append(lista_dispositivos)

    cell = t.cell(1, 21)
    keywords.append(palavras_chave(cell))

    cell = t.cell(1, 22)
    titulo_questionamento_limpo, questionamento_padronizado, tese_padronizada = questionamentos(cell)
    titulo.append(titulo_questionamento_limpo)
    questionamento.append(questionamento_padronizado)
    tese.append(tese_padronizada)

#Criando planilha e acrescentado todas as informações
planilha(numero_processo, processo_anexado, lista_prestadora, lista_cnpj, lista_servico, lista_orgao, lista_vara, processo_principal,processo_associado, tipologia_processo, data_atuacao, data_liminar, doc_judicializacao, status_processo_administrativo, status_processo, data_consulta,regulamento_oficial, dispositivo_oficial, descricao_infracao, titulo, questionamento, tese, keywords, data_primeira_instancia, primeira_inclinacao)